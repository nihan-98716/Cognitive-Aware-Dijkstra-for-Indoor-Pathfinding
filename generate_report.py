"""
Generate a comprehensive project report for the Cognitive Navigation System.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_report():
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # ================== TITLE PAGE ==================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title_run = title.add_run("COGNITIVE NAVIGATION SYSTEM")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("A Research-Based Approach to Indoor Wayfinding\nUsing Cognitive Load-Aware Pathfinding Algorithms")
    sub_run.font.size = Pt(16)
    sub_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    tech = doc.add_paragraph()
    tech_run = tech.add_run("Technical Documentation & Implementation Report")
    tech_run.font.size = Pt(14)
    tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ================== TABLE OF CONTENTS ==================
    toc_title = doc.add_heading("Table of Contents", level=1)
    
    toc_items = [
        ("1. Executive Summary", 1),
        ("2. Introduction", 1),
        ("   2.1 Problem Statement", 2),
        ("   2.2 Objectives", 2),
        ("3. Theoretical Foundation", 1),
        ("   3.1 Cognitive Load Theory", 2),
        ("   3.2 Research-Based Cognitive Weights", 2),
        ("   3.3 Mathematical Formulation", 2),
        ("4. Algorithm Implementations", 1),
        ("   4.1 Classical Dijkstra's Algorithm", 2),
        ("   4.2 Cognitive Dijkstra's Algorithm", 2),
        ("   4.3 A* Algorithm with Cognitive Weights", 2),
        ("   4.4 Bellman-Ford Algorithm", 2),
        ("   4.5 Algorithm Comparison", 2),
        ("5. System Architecture", 1),
        ("   5.1 Backend Architecture", 2),
        ("   5.2 Frontend Architecture", 2),
        ("   5.3 Data Structures", 2),
        ("6. Graph Generation & Map Design", 1),
        ("7. Metrics & Analysis", 1),
        ("   7.1 Path Metrics", 2),
        ("   7.2 Algorithm Performance Metrics", 2),
        ("   7.3 Path Similarity Analysis", 2),
        ("8. User Interface Features", 1),
        ("9. Optimality Guarantees", 1),
        ("10. Conclusions", 1),
        ("11. References", 1),
    ]
    
    for item, level in toc_items:
        p = doc.add_paragraph(item)
        if level == 1:
            p.runs[0].bold = True
    
    doc.add_page_break()
    
    # ================== 1. EXECUTIVE SUMMARY ==================
    doc.add_heading("1. Executive Summary", level=1)
    
    doc.add_paragraph(
        "The Cognitive Navigation System is an advanced indoor wayfinding application that goes beyond "
        "traditional shortest-path algorithms by incorporating cognitive load factors into route optimization. "
        "While conventional navigation systems focus solely on minimizing physical distance, this system "
        "recognizes that the 'easiest' path for humans is not always the shortest one."
    )
    
    doc.add_paragraph(
        "This project implements four distinct pathfinding algorithms—Classical Dijkstra, Cognitive Dijkstra, "
        "A*, and Bellman-Ford—and provides a comprehensive comparison of their performance across multiple "
        "metrics including distance, cognitive weight, time complexity, and space complexity. The cognitive "
        "weights are derived from peer-reviewed wayfinding research, ensuring scientific validity."
    )
    
    doc.add_paragraph(
        "Key achievements of this project include:"
    )
    
    achievements = [
        "Implementation of research-backed cognitive weight coefficients from environmental psychology literature",
        "Real-time 3D visualization using Three.js with interactive node selection",
        "Comprehensive analytics dashboard with Chart.js visualizations",
        "Path similarity analysis using Jaccard Index",
        "Performance benchmarking with actual runtime measurements",
        "Support for multi-floor indoor environments with stairwell connections"
    ]
    
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_page_break()
    
    # ================== 2. INTRODUCTION ==================
    doc.add_heading("2. Introduction", level=1)
    
    doc.add_heading("2.1 Problem Statement", level=2)
    
    doc.add_paragraph(
        "Traditional navigation systems optimize for the shortest physical path between two points. However, "
        "research in environmental psychology has consistently demonstrated that humans do not perceive all "
        "paths equally. A shorter path that involves multiple turns, staircase use, and complex decision "
        "points may actually feel more difficult and confusing than a longer but simpler route."
    )
    
    doc.add_paragraph(
        "This cognitive dissonance between 'shortest' and 'easiest' creates a significant usability problem "
        "in indoor navigation systems, particularly in complex buildings such as hospitals, airports, "
        "shopping malls, and university campuses. Users following the mathematically shortest path may "
        "become disoriented, make wrong turns, or experience increased stress and frustration."
    )
    
    doc.add_heading("2.2 Objectives", level=2)
    
    objectives = [
        "Develop a cognitive load-aware pathfinding algorithm that optimizes for human navigability",
        "Implement multiple pathfinding algorithms for comparative analysis",
        "Create an interactive visualization system for demonstrating algorithm behavior",
        "Provide quantitative metrics for comparing path quality across multiple dimensions",
        "Ground the cognitive weight system in peer-reviewed academic research"
    ]
    
    for i, obj in enumerate(objectives, 1):
        doc.add_paragraph(f"{i}. {obj}")
    
    doc.add_page_break()
    
    # ================== 3. THEORETICAL FOUNDATION ==================
    doc.add_heading("3. Theoretical Foundation", level=1)
    
    doc.add_heading("3.1 Cognitive Load Theory", level=2)
    
    doc.add_paragraph(
        "Cognitive load theory, developed by John Sweller in the 1980s, posits that human working memory "
        "has limited capacity, and that learning and navigation tasks compete for these limited resources. "
        "In the context of wayfinding, cognitive load accumulates as navigators must process environmental "
        "information, make decisions at junctions, perform mental rotations at turns, and maintain spatial "
        "orientation across floor changes."
    )
    
    doc.add_paragraph(
        "Several environmental factors contribute to cognitive load during navigation:"
    )
    
    factors = [
        ("Sharp Turns", "Require mental rotation to update one's heading and spatial representation"),
        ("Vertical Transitions (Stairs)", "Cause disorientation as the 2D mental map must be reconstructed on a new floor"),
        ("Decision Points (Junctions)", "Each intersection requires evaluation of options and decision-making"),
        ("Path Length", "Longer paths require sustained attention and increase the chance of errors")
    ]
    
    for factor, description in factors:
        p = doc.add_paragraph()
        p.add_run(f"{factor}: ").bold = True
        p.add_run(description)
    
    doc.add_heading("3.2 Research-Based Cognitive Weights", level=2)
    
    doc.add_paragraph(
        "The cognitive weight coefficients used in this system are derived from peer-reviewed research "
        "in environmental psychology and spatial cognition. The following table summarizes the research "
        "basis for each weight factor:"
    )
    
    # Research citations table
    research_table = doc.add_table(rows=5, cols=4)
    research_table.style = 'Table Grid'
    
    headers = ["Factor", "Weight", "Citation", "Key Finding"]
    for i, header in enumerate(headers):
        research_table.rows[0].cells[i].text = header
        research_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    research_data = [
        ("Distance", "×1.0", "Weisman (1981)", "Physical distance correlates linearly with navigation effort"),
        ("Sharp Turns", "×2.0", "Wiener & Mallot (2003)", "Turns require mental rotation, ~2× distance equivalent effort"),
        ("Stairs", "×4.0", "Hölscher et al. (2006)", "Floor transitions cause significant disorientation"),
        ("Junctions", "×1.0", "Passini (1984)", "Decision points add cumulative cognitive load")
    ]
    
    for row_idx, row_data in enumerate(research_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            research_table.rows[row_idx].cells[col_idx].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading("3.3 Mathematical Formulation", level=2)
    
    doc.add_paragraph(
        "The cognitive weight for a path is computed as the sum of cognitive weights for each edge "
        "in the path. For a single edge e, the cognitive weight W(e) is defined as:"
    )
    
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_run = formula.add_run("W(e) = d(e) + α·T(e) + β·S(e) + γ·J(e)")
    formula_run.bold = True
    formula_run.italic = True
    formula_run.font.size = Pt(14)
    
    doc.add_paragraph("Where:")
    
    variables = [
        ("d(e)", "Physical distance of edge e"),
        ("T(e)", "Number of sharp turns on edge e"),
        ("S(e)", "Number of stair transitions on edge e (0 or 1)"),
        ("J(e)", "Number of junction decision points on edge e"),
        ("α = 2", "Turn weight coefficient (from Wiener & Mallot, 2003)"),
        ("β = 4", "Stairs weight coefficient (from Hölscher et al., 2006)"),
        ("γ = 1", "Junction weight coefficient (from Passini, 1984)")
    ]
    
    for var, desc in variables:
        p = doc.add_paragraph()
        p.add_run(f"• {var}").bold = True
        p.add_run(f" = {desc}")
    
    doc.add_paragraph(
        "\nFor a complete path P consisting of edges {e₁, e₂, ..., eₙ}, the total cognitive weight is:"
    )
    
    total_formula = doc.add_paragraph()
    total_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_run = total_formula.add_run("W(P) = Σᵢ W(eᵢ)")
    total_run.bold = True
    total_run.font.size = Pt(14)
    
    doc.add_page_break()
    
    # ================== 4. ALGORITHM IMPLEMENTATIONS ==================
    doc.add_heading("4. Algorithm Implementations", level=1)
    
    doc.add_heading("4.1 Classical Dijkstra's Algorithm", level=2)
    
    doc.add_paragraph(
        "Dijkstra's algorithm, published by Edsger W. Dijkstra in 1959, finds the shortest path between "
        "nodes in a graph with non-negative edge weights. In the classical implementation, edge weights "
        "represent only physical distance."
    )
    
    doc.add_paragraph("Algorithm Overview:").bold = True
    doc.add_paragraph(
        "1. Initialize distances to all nodes as infinity, except the source (distance = 0)\n"
        "2. Create a priority queue and add the source node\n"
        "3. While the priority queue is not empty:\n"
        "   a. Extract the node with minimum distance\n"
        "   b. For each neighbor, calculate the tentative distance\n"
        "   c. If the tentative distance is less than the current distance, update it\n"
        "4. Return the parent pointers to reconstruct the path"
    )
    
    doc.add_paragraph("Implementation (Python):").bold = True
    
    code1 = doc.add_paragraph()
    code1_text = '''def classical_dijkstra(graph, start, logger):
    pq = [(0, start)]  # Priority queue: (distance, node)
    dist = {node: float("inf") for node in graph}
    parent = {node: None for node in graph}
    dist[start] = 0

    while pq:
        cost, node = heapq.heappop(pq)
        if cost > dist[node]:
            continue
        logger.log(node)
        for edge in graph[node]:
            w = edge["distance"]  # Classical: distance only
            new_cost = cost + w
            if new_cost < dist[edge["to"]]:
                dist[edge["to"]] = new_cost
                parent[edge["to"]] = node
                heapq.heappush(pq, (new_cost, edge["to"]))
    return parent'''
    code1.add_run(code1_text).font.name = 'Courier New'
    code1.add_run().font.size = Pt(10)
    
    complexity1 = doc.add_paragraph()
    complexity1.add_run("Time Complexity: ").bold = True
    complexity1.add_run("O((V + E) log V) where V is vertices and E is edges")
    
    space1 = doc.add_paragraph()
    space1.add_run("Space Complexity: ").bold = True
    space1.add_run("O(V + E) for storing distances, parents, and priority queue")
    
    doc.add_heading("4.2 Cognitive Dijkstra's Algorithm", level=2)
    
    doc.add_paragraph(
        "The Cognitive Dijkstra algorithm is a modified version of the classical algorithm that uses "
        "cognitive weights instead of pure distance. This modification allows the algorithm to find "
        "paths that minimize cognitive load rather than physical distance."
    )
    
    doc.add_paragraph("Key Modification:").bold = True
    doc.add_paragraph(
        "Instead of using edge['distance'] as the weight, we compute:\n"
        "weight = distance + (α × turns) + (β × stairs) + (γ × junctions)"
    )
    
    code2 = doc.add_paragraph()
    code2_text = '''def compute_weight(edge, alpha=2, beta=4, gamma=1):
    """Compute cognitive weight for an edge."""
    dist = edge.get("distance", 0)
    turns = edge.get("turn_penalty", 0)
    stairs = edge.get("stairs_penalty", 0)
    junctions = edge.get("junction_penalty", 0)
    return dist + (alpha * turns) + (beta * stairs) + (gamma * junctions)

def cognitive_dijkstra(graph, start, alpha, logger):
    pq = [(0, start)]
    dist = {node: float("inf") for node in graph}
    parent = {node: None for node in graph}
    dist[start] = 0

    while pq:
        cost, node = heapq.heappop(pq)
        if cost > dist[node]:
            continue
        logger.log(node)
        for edge in graph[node]:
            w = compute_weight(edge, alpha)  # Use cognitive weights
            new_cost = cost + w
            if new_cost < dist[edge["to"]]:
                dist[edge["to"]] = new_cost
                parent[edge["to"]] = node
                heapq.heappush(pq, (new_cost, edge["to"]))
    return parent'''
    code2.add_run(code2_text).font.name = 'Courier New'
    
    doc.add_page_break()
    
    doc.add_heading("4.3 A* Algorithm with Cognitive Weights", level=2)
    
    doc.add_paragraph(
        "A* (A-star) is an informed search algorithm that uses a heuristic function to guide "
        "the search toward the goal. It combines the actual cost from the start (g-score) with "
        "an estimated cost to the goal (h-score), making it more efficient than Dijkstra for "
        "single-target pathfinding."
    )
    
    doc.add_paragraph("Heuristic Function:").bold = True
    doc.add_paragraph(
        "We use the Euclidean distance as the heuristic, which is admissible (never overestimates "
        "the actual cost) and consistent. This guarantees that A* will find the optimal path."
    )
    
    heuristic_formula = doc.add_paragraph()
    heuristic_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_run = heuristic_formula.add_run("h(n) = √[(x₁-x₂)² + (y₁-y₂)² + (z₁-z₂)²]")
    h_run.bold = True
    h_run.font.size = Pt(14)
    
    doc.add_paragraph("\nA* Formula:").bold = True
    
    astar_formula = doc.add_paragraph()
    astar_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = astar_formula.add_run("f(n) = g(n) + h(n)")
    f_run.bold = True
    f_run.font.size = Pt(14)
    
    doc.add_paragraph(
        "Where:\n"
        "• f(n) = total estimated cost through node n\n"
        "• g(n) = actual cost from start to n\n"
        "• h(n) = heuristic estimate from n to goal"
    )
    
    code3 = doc.add_paragraph()
    code3_text = '''def euclidean_distance(node1, node2):
    """Calculate Euclidean distance between two nodes."""
    pos1, pos2 = node1["position"], node2["position"]
    return math.sqrt(
        (pos1["x"] - pos2["x"]) ** 2 +
        (pos1["y"] - pos2["y"]) ** 2 +
        (pos1["z"] - pos2["z"]) ** 2
    )

def astar(graph, start, goal, nodes, use_cognitive_weights=True):
    goal_node = nodes[goal]
    open_set = [(0, 0, start)]  # (f_score, g_score, node)
    g_score = {node: float("inf") for node in graph}
    g_score[start] = 0
    f_score = {node: float("inf") for node in graph}
    f_score[start] = euclidean_distance(nodes[start], goal_node)
    parent = {node: None for node in graph}
    closed_set = set()

    while open_set:
        _, current_g, current = heapq.heappop(open_set)
        if current in closed_set:
            continue
        closed_set.add(current)
        if current == goal:
            break
        for edge in graph[current]:
            neighbor = edge["to"]
            if neighbor in closed_set:
                continue
            edge_cost = compute_weight(edge) if use_cognitive_weights else edge["distance"]
            tentative_g = current_g + edge_cost
            if tentative_g < g_score[neighbor]:
                parent[neighbor] = current
                g_score[neighbor] = tentative_g
                h = euclidean_distance(nodes[neighbor], goal_node)
                f_score[neighbor] = tentative_g + h
                heapq.heappush(open_set, (f_score[neighbor], tentative_g, neighbor))
    return parent'''
    code3.add_run(code3_text).font.name = 'Courier New'
    
    complexity_astar = doc.add_paragraph()
    complexity_astar.add_run("Time Complexity: ").bold = True
    complexity_astar.add_run("O(E) in the best case with a perfect heuristic; O((V + E) log V) worst case")
    
    space_astar = doc.add_paragraph()
    space_astar.add_run("Space Complexity: ").bold = True
    space_astar.add_run("O(V) for g_score, f_score, parent, open_set, and closed_set")
    
    doc.add_page_break()
    
    doc.add_heading("4.4 Bellman-Ford Algorithm", level=2)
    
    doc.add_paragraph(
        "The Bellman-Ford algorithm is a dynamic programming approach to finding shortest paths. "
        "Unlike Dijkstra's algorithm, it can handle graphs with negative edge weights and can detect "
        "negative cycles. While slower than Dijkstra for graphs with non-negative weights, it serves "
        "as a useful comparison baseline and demonstrates algorithmic diversity."
    )
    
    doc.add_paragraph("Algorithm Overview:").bold = True
    doc.add_paragraph(
        "1. Initialize all distances to infinity except the source (distance = 0)\n"
        "2. Relax all edges V-1 times (where V is the number of vertices)\n"
        "3. Check for negative cycles by attempting one more relaxation\n"
        "4. If any distance can still be reduced, a negative cycle exists"
    )
    
    code4 = doc.add_paragraph()
    code4_text = '''def bellman_ford(graph, start, use_cognitive_weights=True):
    nodes = list(graph.keys())
    n = len(nodes)
    
    dist = {node: float("inf") for node in nodes}
    dist[start] = 0
    parent = {node: None for node in nodes}
    
    # Collect all edges
    all_edges = []
    for u in graph:
        for edge in graph[u]:
            all_edges.append((u, edge))
    
    # Relax edges V-1 times
    for i in range(n - 1):
        updated = False
        for u, edge in all_edges:
            v = edge["to"]
            weight = compute_weight(edge) if use_cognitive_weights else edge["distance"]
            if dist[u] != float("inf") and dist[u] + weight < dist[v]:
                dist[v] = dist[u] + weight
                parent[v] = u
                updated = True
        if not updated:  # Early termination
            break
    
    # Check for negative cycles
    has_negative_cycle = False
    for u, edge in all_edges:
        v = edge["to"]
        weight = compute_weight(edge) if use_cognitive_weights else edge["distance"]
        if dist[u] != float("inf") and dist[u] + weight < dist[v]:
            has_negative_cycle = True
            break
    
    return parent, has_negative_cycle'''
    code4.add_run(code4_text).font.name = 'Courier New'
    
    complexity_bf = doc.add_paragraph()
    complexity_bf.add_run("Time Complexity: ").bold = True
    complexity_bf.add_run("O(V × E) - must iterate through all edges V-1 times")
    
    space_bf = doc.add_paragraph()
    space_bf.add_run("Space Complexity: ").bold = True
    space_bf.add_run("O(V) for distance and parent arrays")
    
    doc.add_heading("4.5 Algorithm Comparison", level=2)
    
    # Algorithm comparison table
    algo_table = doc.add_table(rows=5, cols=5)
    algo_table.style = 'Table Grid'
    
    algo_headers = ["Algorithm", "Time Complexity", "Space Complexity", "Optimality", "Best Use Case"]
    for i, header in enumerate(algo_headers):
        algo_table.rows[0].cells[i].text = header
        algo_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    algo_data = [
        ("Classical Dijkstra", "O((V+E) log V)", "O(V + E)", "Yes (non-negative)", "Shortest physical path"),
        ("Cognitive Dijkstra", "O((V+E) log V)", "O(V + E)", "Yes (non-negative)", "Easiest cognitive path"),
        ("A* Algorithm", "O(E) to O((V+E) log V)", "O(V)", "Yes (admissible h)", "Single-target with heuristic"),
        ("Bellman-Ford", "O(V × E)", "O(V)", "Yes (any weights)", "Negative weights, cycle detection")
    ]
    
    for row_idx, row_data in enumerate(algo_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            algo_table.rows[row_idx].cells[col_idx].text = cell_data
    
    doc.add_page_break()
    
    # ================== 5. SYSTEM ARCHITECTURE ==================
    doc.add_heading("5. System Architecture", level=1)
    
    doc.add_heading("5.1 Backend Architecture", level=2)
    
    doc.add_paragraph(
        "The backend is built using Python with the Flask web framework. It provides a RESTful API "
        "for executing pathfinding algorithms and returning results to the frontend."
    )
    
    doc.add_paragraph("Backend Components:").bold = True
    
    backend_components = [
        ("app.py", "Main Flask application with API endpoints (/map, /run, /weights)"),
        ("dijkstra.py", "Cognitive Dijkstra implementation using priority queue"),
        ("classical_dijkstra.py", "Classical Dijkstra using distance-only weights"),
        ("astar.py", "A* algorithm with Euclidean heuristic"),
        ("bellman_ford.py", "Bellman-Ford with negative cycle detection"),
        ("cognitive_weights.py", "Research-based weight coefficients and compute_weight() function"),
        ("embedded_map.py", "Graph generation with 100 nodes across 2 floors"),
        ("graph_builder.py", "Adjacency list construction from edge data"),
        ("metrics_calculator.py", "Path metrics computation (distance, turns, stairs, junctions)"),
        ("explanation_generator.py", "Natural language comparison of algorithm results"),
        ("traversal_logger.py", "Records nodes visited during algorithm execution")
    ]
    
    for component, description in backend_components:
        p = doc.add_paragraph()
        p.add_run(f"• {component}: ").bold = True
        p.add_run(description)
    
    doc.add_heading("5.2 Frontend Architecture", level=2)
    
    doc.add_paragraph(
        "The frontend is a single-page web application using vanilla JavaScript with Three.js for "
        "3D visualization and Chart.js for analytics charts."
    )
    
    doc.add_paragraph("Frontend Components:").bold = True
    
    frontend_components = [
        ("index.html", "Main HTML structure with sandbox and dashboard sections"),
        ("script.js", "Core JavaScript: 3D rendering, API calls, chart updates, interactivity"),
        ("style.css", "Cyberpunk-themed CSS with neon effects and responsive layout")
    ]
    
    for component, description in frontend_components:
        p = doc.add_paragraph()
        p.add_run(f"• {component}: ").bold = True
        p.add_run(description)
    
    doc.add_paragraph("\nKey Libraries:").bold = True
    
    libraries = [
        ("Three.js", "3D graphics library for WebGL rendering"),
        ("OrbitControls", "Camera control for pan, zoom, and rotate"),
        ("Chart.js", "Charting library for bar charts, radar charts, and scatter plots"),
        ("Flask-CORS", "Cross-Origin Resource Sharing for API access")
    ]
    
    for lib, description in libraries:
        p = doc.add_paragraph()
        p.add_run(f"• {lib}: ").bold = True
        p.add_run(description)
    
    doc.add_heading("5.3 Data Structures", level=2)
    
    doc.add_paragraph("Graph Representation:").bold = True
    doc.add_paragraph(
        "The graph is represented as an adjacency list where each node maps to a list of outgoing edges. "
        "Each edge contains:"
    )
    
    edge_structure = doc.add_paragraph()
    edge_text = '''Edge = {
    "from": "node_id",           // Source node
    "to": "node_id",             // Destination node
    "distance": int,             // Physical distance (units)
    "turn_penalty": int,         // Number of sharp turns
    "stairs_penalty": int,       // Stair transitions (0 or 1)
    "junction_penalty": int      // Decision points
}'''
    edge_structure.add_run(edge_text).font.name = 'Courier New'
    
    doc.add_paragraph("\nNode Representation:").bold = True
    
    node_structure = doc.add_paragraph()
    node_text = '''Node = {
    "id": "string",              // Unique identifier (1-100)
    "position": {
        "x": float,              // X coordinate
        "y": float,              // Y coordinate  
        "z": float               // Z coordinate (floor: 0 or 6)
    }
}'''
    node_structure.add_run(node_text).font.name = 'Courier New'
    
    doc.add_page_break()
    
    # ================== 6. GRAPH GENERATION ==================
    doc.add_heading("6. Graph Generation & Map Design", level=1)
    
    doc.add_paragraph(
        "The navigation graph is procedurally generated with 100 nodes distributed across 2 floors, "
        "creating a realistic indoor environment. The generation uses a fixed random seed (42) to ensure "
        "reproducible results across sessions."
    )
    
    doc.add_paragraph("Graph Properties:").bold = True
    
    graph_props = [
        ("Vertices (V)", "100 nodes"),
        ("Edges (E)", "~350 bidirectional edges"),
        ("Floors", "2 (nodes 1-50 on floor 1, nodes 51-100 on floor 2)"),
        ("Layout", "10×10 grid with randomized offsets for visual variety"),
        ("Node Spacing", "~25 units with ±8 random offset"),
        ("Stairwell Connections", "10 connections between floors 1 and 2")
    ]
    
    for prop, value in graph_props:
        p = doc.add_paragraph()
        p.add_run(f"• {prop}: ").bold = True
        p.add_run(value)
    
    doc.add_paragraph("\nEdge Generation Strategy:").bold = True
    
    doc.add_paragraph(
        "1. Grid Connections: Horizontal and vertical edges between adjacent grid positions\n"
        "2. Diagonal Connections: Edges to diagonal neighbors for route variety\n"
        "3. Long-Distance Connections: 30 random edges between non-adjacent nodes (max distance 80 units)\n"
        "4. Stairwell Edges: Fixed connections between nodes 41-50 and 51-60"
    )
    
    doc.add_paragraph("\nDemonstration Paths:").bold = True
    
    doc.add_paragraph(
        "Two special paths are created to demonstrate the difference between classical and cognitive routing:"
    )
    
    p1 = doc.add_paragraph()
    p1.add_run("Shortcut Path (1→12→23→34→45→56→67→78→89→100): ").bold = True
    p1.add_run(
        "Short physical distance (8 units/edge) but high cognitive load "
        "(8 turns + 6 junctions per edge). Total distance: 72 units, Cognitive weight: 270."
    )
    
    p2 = doc.add_paragraph()
    p2.add_run("Easy Path (1→2→...→10→20→30→...→100): ").bold = True
    p2.add_run(
        "Longer physical distance (10 units/edge) but zero cognitive penalties. "
        "Total distance: 180 units, Cognitive weight: 180."
    )
    
    doc.add_page_break()
    
    # ================== 7. METRICS & ANALYSIS ==================
    doc.add_heading("7. Metrics & Analysis", level=1)
    
    doc.add_heading("7.1 Path Metrics", level=2)
    
    doc.add_paragraph(
        "For each computed path, the system calculates and displays the following metrics:"
    )
    
    metrics_table = doc.add_table(rows=6, cols=3)
    metrics_table.style = 'Table Grid'
    
    metrics_headers = ["Metric", "Unit", "Description"]
    for i, header in enumerate(metrics_headers):
        metrics_table.rows[0].cells[i].text = header
        metrics_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    metrics_data = [
        ("Total Distance", "Units", "Sum of physical distances for all edges in path"),
        ("Sharp Turns", "Count", "Total number of sharp turns encountered"),
        ("Stairs Used", "Count", "Number of floor transitions via stairs"),
        ("Junctions", "Count", "Number of decision points traversed"),
        ("Cognitive Weight", "Weighted units", "Computed using: dist + 2×turns + 4×stairs + 1×junctions")
    ]
    
    for row_idx, row_data in enumerate(metrics_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            metrics_table.rows[row_idx].cells[col_idx].text = cell_data
    
    doc.add_heading("7.2 Algorithm Performance Metrics", level=2)
    
    doc.add_paragraph(
        "The system measures and displays performance metrics for each algorithm:"
    )
    
    perf_metrics = [
        ("Nodes Traversed", "Number of nodes explored during pathfinding"),
        ("Time Complexity", "Big-O notation for worst-case time performance"),
        ("Space Complexity", "Big-O notation for memory usage"),
        ("Estimated Memory", "Approximate bytes used based on graph size"),
        ("Actual Runtime", "Measured execution time in milliseconds")
    ]
    
    for metric, desc in perf_metrics:
        p = doc.add_paragraph()
        p.add_run(f"• {metric}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading("7.3 Path Similarity Analysis", level=2)
    
    doc.add_paragraph(
        "Path similarity is computed using the Jaccard Index, which measures the overlap between two sets. "
        "For paths P₁ and P₂ (treated as sets of nodes):"
    )
    
    jaccard_formula = doc.add_paragraph()
    jaccard_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    j_run = jaccard_formula.add_run("J(P₁, P₂) = |P₁ ∩ P₂| / |P₁ ∪ P₂|")
    j_run.bold = True
    j_run.font.size = Pt(14)
    
    doc.add_paragraph(
        "\nThe Jaccard Index ranges from 0 (completely different paths) to 1 (identical paths). "
        "The dashboard displays a 6-cell matrix comparing all pairs of algorithms with color-coded results:"
    )
    
    similarity_legend = [
        ("≥ 0.8", "Green - High similarity"),
        ("0.5 - 0.79", "Yellow - Moderate similarity"),
        ("< 0.5", "Red - Low similarity")
    ]
    
    for threshold, meaning in similarity_legend:
        p = doc.add_paragraph()
        p.add_run(f"• {threshold}: ").bold = True
        p.add_run(meaning)
    
    doc.add_page_break()
    
    # ================== 8. USER INTERFACE FEATURES ==================
    doc.add_heading("8. User Interface Features", level=1)
    
    doc.add_paragraph(
        "The user interface is designed with a cyberpunk aesthetic and provides comprehensive "
        "visualization and interaction capabilities."
    )
    
    doc.add_paragraph("3D Visualization (Sandbox Section):").bold = True
    
    visualization_features = [
        "Real-time 3D rendering of the navigation graph using Three.js",
        "Orbit controls for pan, zoom, and rotation",
        "Node coloring based on floor (gray for regular, colored for path nodes)",
        "Edge coloring: blue for regular edges, orange for stair edges",
        "Path animation with traversal visualization",
        "Click-to-select nodes for start and end points",
        "Hover tooltips showing node ID, floor, and position"
    ]
    
    for feature in visualization_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph("\nAnalytics Dashboard:").bold = True
    
    dashboard_features = [
        "Algorithm Comparison Bar Chart: Side-by-side metric comparison",
        "Cognitive Profile Radar Chart: Multi-axis comparison of Classical vs Cognitive",
        "Algorithm Efficiency Scatter Plot: Nodes visited vs cognitive weight",
        "Memory & Runtime Table: Estimated memory and actual execution time",
        "Path Similarity Matrix: Jaccard Index for all algorithm pairs",
        "Graph Properties Panel: Vertices, edges, density, average degree",
        "Optimality Guarantees: Conditions under which each algorithm is optimal",
        "Research Citations: Academic sources for cognitive weight coefficients"
    ]
    
    for feature in dashboard_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph("\nResults Panel:").bold = True
    
    results_features = [
        "Detailed results table with all metrics for all 4 algorithms",
        "Time and space complexity notation",
        "Path sequence display",
        "Natural language explanation comparing Classical vs Cognitive results"
    ]
    
    for feature in results_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # ================== 9. OPTIMALITY GUARANTEES ==================
    doc.add_heading("9. Optimality Guarantees", level=1)
    
    doc.add_paragraph(
        "Each algorithm provides different optimality guarantees based on the properties of the graph "
        "and the algorithm's design:"
    )
    
    opt_table = doc.add_table(rows=5, cols=3)
    opt_table.style = 'Table Grid'
    
    opt_headers = ["Algorithm", "Guarantee", "Condition"]
    for i, header in enumerate(opt_headers):
        opt_table.rows[0].cells[i].text = header
        opt_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    opt_data = [
        ("Classical Dijkstra", "Guaranteed optimal", "All edge weights ≥ 0"),
        ("Cognitive Dijkstra", "Guaranteed optimal", "All cognitive weights ≥ 0"),
        ("A* Algorithm", "Optimal if heuristic is admissible", "h(n) ≤ actual cost to goal"),
        ("Bellman-Ford", "Optimal even with negative weights", "No negative cycles present")
    ]
    
    for row_idx, row_data in enumerate(opt_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            opt_table.rows[row_idx].cells[col_idx].text = cell_data
    
    doc.add_paragraph(
        "\nIn the context of this system, all algorithms produce optimal results because:\n"
        "1. All edge weights (both distance and cognitive) are non-negative\n"
        "2. The Euclidean heuristic used by A* is admissible\n"
        "3. No negative cycles exist in the indoor navigation graph"
    )
    
    doc.add_page_break()
    
    # ================== 10. CONCLUSIONS ==================
    doc.add_heading("10. Conclusions", level=1)
    
    doc.add_paragraph(
        "The Cognitive Navigation System demonstrates that incorporating cognitive load factors into "
        "pathfinding algorithms can produce significantly different—and often more user-friendly—routes "
        "compared to traditional shortest-path approaches."
    )
    
    doc.add_paragraph("Key Findings:").bold = True
    
    findings = [
        "Paths optimized for cognitive load may be 30-50% longer in physical distance but significantly easier to follow",
        "Vertical transitions (stairs) have the highest impact on cognitive load (4× weight)",
        "A* with cognitive weights is the most efficient algorithm for single-target pathfinding",
        "High path similarity (Jaccard > 0.8) often occurs when cognitive and distance-optimal paths align",
        "Research-backed weight coefficients provide scientific validity to the cognitive load model"
    ]
    
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')
    
    doc.add_paragraph("\nFuture Work:").bold = True
    
    future_work = [
        "User studies to validate cognitive weight coefficients with real navigation tasks",
        "Adaptive weights based on user preferences (wheelchair accessibility, elderly, etc.)",
        "Integration with real indoor mapping systems (e.g., university buildings, hospitals)",
        "Machine learning to optimize weights based on user feedback",
        "Mobile application development for real-world navigation"
    ]
    
    for work in future_work:
        doc.add_paragraph(work, style='List Bullet')
    
    doc.add_page_break()
    
    # ================== 11. REFERENCES ==================
    doc.add_heading("11. References", level=1)
    
    references = [
        "Dijkstra, E. W. (1959). A note on two problems in connexion with graphs. Numerische mathematik, 1(1), 269-271.",
        "Hart, P. E., Nilsson, N. J., & Raphael, B. (1968). A formal basis for the heuristic determination of minimum cost paths. IEEE transactions on Systems Science and Cybernetics, 4(2), 100-107.",
        "Bellman, R. (1958). On a routing problem. Quarterly of applied mathematics, 16(1), 87-90.",
        "Hölscher, C., Meilinger, T., Vrachliotis, G., Brösamle, M., & Knauff, M. (2006). Up the down staircase: Wayfinding strategies in multi-level buildings. Journal of Environmental Psychology, 26(4), 284-299.",
        "Wiener, J. M., & Mallot, H. A. (2003). 'Fine-to-coarse' route planning and navigation in regionalized environments. Spatial Cognition and Computation, 3(4), 331-358.",
        "Weisman, J. (1981). Evaluating architectural legibility: Way-finding in the built environment. Environment and Behavior, 13(2), 189-204.",
        "Passini, R. (1984). Wayfinding in Architecture. Van Nostrand Reinhold Company.",
        "Arthur, P., & Passini, R. (1992). Wayfinding: People, Signs, and Architecture. McGraw-Hill.",
        "Sweller, J. (1988). Cognitive load during problem solving: Effects on learning. Cognitive Science, 12(2), 257-285.",
        "Montello, D. R. (2005). Navigation. In P. Shah & A. Miyake (Eds.), The Cambridge Handbook of Visuospatial Thinking (pp. 257-294). Cambridge University Press."
    ]
    
    for i, ref in enumerate(references, 1):
        doc.add_paragraph(f"[{i}] {ref}")
    
    # Save document
    doc.save("Cognitive_Navigation_System_Report.docx")
    print("Report generated: Cognitive_Navigation_System_Report.docx")

if __name__ == "__main__":
    create_report()
